from .models import Department, Schedule, College
import docx
from docx.shared import Mm, Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from django.templatetags.static import static
import os
from pathlib import Path
import comtypes.client



def tablegenerator(college_id, user_id, some_list, batch):
    doc = docx.Document()
    section = doc.sections[0]
    section.page_height = Mm(279.4)
    section.page_width = Mm(471.8)
    section.orientation = WD_ORIENT.LANDSCAPE
    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(20)
    font_object.name = 'Times New Roman'
    days = {1:"Monday", 2:"Tuesday", 3:"Wednesday", 4:"Thursday", 5:"Friday"}
    times = {0:"PERIOD",1:"8:00 - 8:55", 2:"9:00 - 9:55", 3:"10:30 - 11:25", 4:"11:30 - 12:25", 5:"1:00 - 1:55", 6:"2:00 - 2:55", 7:"3:00 - 3:55", 8:"4:00 - 4:55", 9:"5:00 - 5:55", 10:"6:00 - 6:55"}
    departments = Department.objects.filter(college_main_id=college_id)
    max_yg = (departments.order_by("-max_yg").first()).max_yg
    college = College.objects.get(id=college_id)
    rows_per_day = college.rows_per_day
    days_per_week = college.days_per_week
    d_counter = 1
    years = ["FIRST", "SECOND", "THIRD", "FOURTH", "FIFTH", "SIXTH", "SEVENTH", "HEIGTH", "NINTH"]
    for m in range(1, max_yg+1):
        for key, day in list(days.items()):
            header1 = doc.add_heading(f"College Of {college.name}")
            header1.alignment = 1
            par1 = doc.add_heading(f"{years[m-1]} YEAR", 4)
            par1.alignment = 1
            par1.add_run(" ").bold = True
            par2 = doc.add_paragraph(f"{day}")
            par2.alignment = 1
            par2.add_run(".").bold = True
            table = doc.add_table(rows=11, cols=len(departments)+1)
            table.style = 'Table Grid'
            doc.add_paragraph('')
            doc.add_page_break()
            trow = table.rows[0].cells
            for r in table.rows: 
                if r not in trow: r.height = Cm(2.5)
            counter = 1
            for department in departments:
                trow[counter].text = department.name
                counter = counter + 1
            table.rows[0].height = Cm(1.0)
            tcol = table.columns[0].cells
            counter = 0
            for t in list(times.values()):
                tcol[counter].text = t
                counter = counter + 1
            for sch in some_list:
                year_group = sch.year_group
                max_yg_row = year_group * (rows_per_day * days_per_week)
                min_yg_row = (max_yg_row + 1) - (rows_per_day * days_per_week)
                if sch.year_group == m:
                    if sch.day == key:
                        ss_row = (sch.row - min_yg_row + 1) % rows_per_day
                        if ss_row == 0: ss_row = rows_per_day
                        table.rows[ss_row].cells[sch.column].text = (f"\n{sch.course_code}\t\n{sch.location_name }\t\n{sch.lecturer_name}\n")
    file_name = str(college.name) + "_" + str(batch) + ".docx"
    doc.save(file_name)
   