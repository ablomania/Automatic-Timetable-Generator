from .models import Department, Schedule, College
import docx
from docx.shared import Mm, Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from django.templatetags.static import static
import os
from pathlib import Path
import comtypes.client


def examTableGenerator(college_id, user_id, days_list, exams_list, exams, table_id, courses):
    doc = docx.Document()
    section = doc.sections[0]
    section.page_height = Mm(279.4)
    section.page_width = Mm(471.8)
    section.orientation = WD_ORIENT.LANDSCAPE
    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(20)
    font_object.name = 'Times New Roman'
    table_headers = ['TIME', 'DEPARTMENT / YEAR GROUP', 'COURSE CODE', 'LOCATIONS', 'LECTURER', 'CLASS SIZE / STUDENTS PER ROOM']
    for day in list(days_list):
        n_rows = 0
        for x in list(exams_list.values()):
            if x.day == day: n_rows = n_rows + 1
        head1 = doc.add_heading(f"Day {day}")
        table = doc.add_table(rows=n_rows+1, cols=6)
        trow = table.rows[0].cells
        for r in table.rows: 
            if r not in trow: r.height = Cm(2.5)
        counter = 0
        for table_header in list(table_headers):
            index = table_headers.index(table_header)
            trow[index].text = table_header
            counter = counter + 1
            table.rows[0].height = Cm(1.0)
        counter  = 1
        index = 0
        for el in list(exams_list.values()):
            for s_row in range(1, n_rows+1):
                cell = table.rows[s_row].cells
                if el.day == day:
                    cell[0].text = str(el.time)
                

    doc.save("aaa.docx")
        

    
